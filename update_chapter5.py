# -*- coding: utf-8 -*-
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import copy

doc = Document('实验总结报告.docx')

# Find paragraph indices for 五、实验问题与解决方法 section
chapter5_heading_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith('五、实验问题与解决方法'):
        chapter5_heading_idx = i
        break

if chapter5_heading_idx is None:
    print("ERROR: 未找到「五、实验问题与解决方法」")
    sys.exit(1)

print(f"找到第五章标题: 段落[{chapter5_heading_idx}]")

# The placeholder structure:
# [141] 五、实验问题与解决方法  <- heading, keep
# [142] （由用户补充具体的问题描述...）<- replace with intro
# [143] (empty) <- delete
# [144] 问题 1：【待填写】<- replace
# [145] 解决方案：【待填写】<- replace
# [146] (empty) <- delete
# [147] 问题 2：【待填写】<- replace
# [148] 解决方案：【待填写】<- replace
# [149] (empty) <- delete
# [150] 问题 3：【待填写】<- replace
# [151] 解决方案：【待填写】<- replace
# [152] (empty) <- delete
# [153] 六、实验总结与心得体会  <- keep

# The plan: We'll work backwards, deleting paragraphs from 152 down to 143
# Then insert new content after the heading paragraph

# Let's verify the paragraphs we're targeting
print("\n当前第五章占位内容：")
for idx in range(chapter5_heading_idx, min(chapter5_heading_idx + 13, len(doc.paragraphs))):
    text = doc.paragraphs[idx].text[:80] if doc.paragraphs[idx].text else '(空)'
    print(f"  [{idx}] {text}")

# Content to insert after the heading paragraph
intro_text = (
    "本小组在实验过程中遇到了以下三个主要技术问题，经过团队协作与资料查阅均已成功解决。"
)

problems = [
    {
        "title": "问题 1：深度学习模型在 Notebook 上无法正常训练",
        "desc": (
            "问题描述：在华为云 ModelArts Notebook 环境中运行 MindSpore Autoencoder 训练脚本时，"
            "模型训练无法正常启动，多次报错提示设备不兼容或算子不支持。经排查发现，Notebook 实例默认"
            "配置了 Ascend NPU 设备上下文，而实验环境实际不具备 NPU 硬件，导致 MindSpore 框架在尝试"
            "将算子调度到 Ascend 设备时失败。",
            "解决方案：通过 AI 编程助手的辅助分析，了解到可以通过修改 MindSpore 的 context 设置，"
            "将设备目标强制切换为 CPU 模式。具体操作方法：在训练脚本开头调用 "
            "context.set_context(mode=context.PYNATIVE_MODE, device_target=\"CPU\")，"
            "使 MindSpore 在 CPU 上以 PyNative 动态图模式执行。修改后模型训练成功启动，200 epoch 后"
            "损失函数收敛至预期水平，异常检测结果符合数据分布特征。"
        )
    },
    {
        "title": "问题 2：ModelArts 自定义服务部署报错——基类继承错误",
        "desc": (
            "问题描述：将训练好的 Autoencoder 模型部署为 ModelArts 在线推理服务时，编写的 "
            "customize_service.py 部署脚本在上传后持续报错，错误信息提示服务类找不到或初始化失败。"
            "经代码审查发现，My_ModelService 类错误地继承了 model_service.model_service.ModelService "
            "基类，而 ModelArts 自定义模型的服务入口要求继承 SingleNodeService 类，两者接口规范不同，"
            "导致服务框架无法正确识别和初始化推理服务实例。",
            "解决方案：通过查阅华为云 ModelArts 官方文档中关于自定义模型推理代码的规范说明，明确了"
            "单节点部署场景下必须继承 model_service.model_service.SingleNodeService 类，并实现 "
            "_preprocess、_inference、_postprocess 三个核心方法。将 customize_service.py 中的类继承关系"
            "从 ModelService 修正为 SingleNodeService 后重新上传，服务部署成功启动，API 端点正常响应。"
        )
    },
    {
        "title": "问题 3：ModelArts 在线服务部署完成后不知如何调用 API",
        "desc": (
            "问题描述：模型在线服务部署状态显示「运行中」后，小组成员对如何通过 HTTP 请求调用推理服务"
            "缺乏经验。最初尝试直接访问服务端点 URL，但返回 403 认证失败；尝试参考新版本 ModelArts 的"
            "Token 鉴权方式调用，同样未能成功。经对 ModelArts 控制台的「服务详情」页面深入排查后，发现"
            "当前使用的旧版在线部署方式需要特定的认证信息组合。",
            "解决方案：最终确定旧版 ModelArts 在线服务部署的 API 调用需要使用两个关键信息：(1) API "
            "接口公网地址（即服务的「调用URL」，从 ModelArts 控制台「服务详情」->「调用信息」获取）；"
            "(2) AppCode 值（在 API 网关的 App 管理中设置，用于 X-Apig-AppCode 请求头认证）。"
            "调用方式为 HTTP POST 请求，请求头中设置 X-Apig-AppCode: <AppCode值>，请求体中传入 "
            "包含 records 和 fields 的 JSON 数据。按照此方法成功调用了推理 API，返回了包含 "
            "anomaly_count、anomaly_ratio、anomalies 等字段的检测结果 JSON，并在前端的深度学习"
            "结果展示界面中成功渲染。"
        )
    },
]

# Work backwards to delete placeholder paragraphs (index 152 down to 143)
# The paragraphs from 144 to 152 are placeholders we want to replace
# Paragraph 143 is just an empty line we want to delete
# Paragraph 142 is the hint text we want to replace

# We need to access the XML element and remove paragraphs
# Find the body element
body = doc.element.body

# Get paragraph elements from 142 to 152 (inclusive)
paras_to_remove = []
for idx in range(142, 152 + 1):
    if idx < len(doc.paragraphs):
        para_elem = doc.paragraphs[idx]._element
        paras_to_remove.append(para_elem)

# Remove each paragraph element from the body
for elem in paras_to_remove:
    body.remove(elem)

print(f"\n已删除段落 142 到 152 的占位内容")

# Now insert new paragraphs after the heading
# The heading is at index chapter5_heading_idx in the original document
# After deletion, we need to find the heading element and insert after it

# Find the heading paragraph element
from docx.oxml.ns import qn
heading_elem = doc.paragraphs[chapter5_heading_idx]._element

# Helper function to insert a new paragraph after a given paragraph element
def insert_paragraph_after(after_element, text, bold=False, font_size=None, first_line_indent=None):
    """Insert a new paragraph after the given element."""
    new_p = doc.add_paragraph()
    # Move the new paragraph element right after after_element
    after_element.addnext(new_p._element)

    if text:
        run = new_p.add_run(text)
        if bold:
            run.bold = True
        if font_size:
            run.font.size = font_size
        if first_line_indent:
            new_p.paragraph_format.first_line_indent = first_line_indent
    return new_p

# Insert intro paragraph after heading
intro_p = insert_paragraph_after(heading_elem, intro_text)
print("已插入第五章引言段落")

# Insert three problems - each problem has a title and description
# Keep track of the last inserted element for successive insertions
last_elem = intro_p._element

for i, prob in enumerate(problems):
    # Insert problem title
    title_p = insert_paragraph_after(last_elem, prob["title"], bold=True, font_size=Pt(12))
    last_elem = title_p._element

    # Insert problem description
    desc_p = insert_paragraph_after(last_elem, prob["desc"], first_line_indent=Pt(21))
    last_elem = desc_p._element

    # Add empty line between problems (except after the last one)
    if i < len(problems) - 1:
        empty_p = insert_paragraph_after(last_elem, "")
        last_elem = empty_p._element

    print(f"已插入问题 {i+1}：{prob['title']}")

# Save the document
doc.save('实验总结报告.docx')
print("\n文档已保存！第五章内容更新完成。")

# Verify
doc2 = Document('实验总结报告.docx')
print("\n=== 验证：第五章完整内容 ===")
found = False
for i, p in enumerate(doc2.paragraphs):
    if p.text.startswith('五、实验问题'):
        found = True
    if found:
        text = p.text[:120] if p.text else '(空)'
        print(f"[{i}] {text}")
    if found and p.text.startswith('六、实验总结'):
        break
