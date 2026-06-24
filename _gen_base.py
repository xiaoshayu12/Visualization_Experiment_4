# -*- coding: utf-8 -*-
"""
Generate 4 individual task report .docx files for team members.
Usage: pip install python-docx && python gen.py
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- helpers ----------
def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)

def set_run_font(run, name="Microsoft YaHei", size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return h

def add_para(doc, text, size=12, bold=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p

def add_table_with_header(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2F5496")
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D6E4F0")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

def build_report(name, sections, sub_task_zh):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    set_default_style(doc)
    for _ in range(6):
        doc.add_paragraph()

    add_para(doc, "\u53ef\u89c6\u5316\u8bfe\u7a0b\u5b9e\u9a8c\u4e09", size=26, bold=True,
             color=(47, 84, 150), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
    add_para(doc, "\u4e2a\u4eba\u4efb\u52a1\u62a5\u544a", size=22, bold=True,
             color=(47, 84, 150), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    add_para(doc, sub_task_zh, size=18, bold=True,
             color=(0, 0, 0), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

    info_lines = [
        ("\u59d3\u540d", name),
        ("\u65e5\u671f", "2026\u5e746\u670814\u65e5"),
        ("\u8bfe\u7a0b", "\u53ef\u89c6\u5316"),
        ("\u5b9e\u9a8c", "\u5b9e\u9a8c\u4e09: \u5168\u7403\u6d77\u6d0b\u5927\u6c14\u8026\u5408\u65f6\u7a7a\u53ef\u89c6\u5206\u6790"),
    ]
    for label, val in info_lines:
        add_para(doc, label + ":  " + val, size=14,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    doc.add_page_break()

    for sec_title, sec_body in sections:
        add_heading_styled(doc, sec_title, level=1)
        for item in sec_body:
            if item[0] == "para":
                add_para(doc, item[1], size=12, space_after=8)
            elif item[0] == "bold":
                add_para(doc, item[1], size=12, bold=True, space_after=4)
            elif item[0] == "bullet":
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(item[1])
                set_run_font(run, size=11)
            elif item[0] == "table":
                add_table_with_header(doc, item[1], item[2], item[3] if len(item) > 3 else None)

    out_dir = r"C:\Users\DELL\Desktop\文件\可视化\实验三"
    fname = os.path.join(out_dir, "\u4e2a\u4eba\u4efb\u52a1\u62a5\u544a_" + name + ".docx")
    doc.save(fname)
    print("Saved:", fname)
    return fname

# ================================================================
# CONTENT DATA - placed in a separate file for readability
# The content is injected via exec() below
# ================================================================

if __name__ == "__main__":
    print("This bootstrap script writes gen.py. Run gen.py to generate docx files.")
